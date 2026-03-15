import io
from docx import Document

from src.docs import generate_doc


def create_test_template():
    """Создает простой шаблон для тестирования."""
    doc = Document()
    doc.add_paragraph('Статус: {{ status }}')
    doc.add_paragraph('Имя: {{ user.name }}')
    doc.add_paragraph('Фамилия: {{ user.surname }}')
    doc.add_paragraph('Отчество: {{ user.father_name }}')
    doc.add_paragraph('Роль: {{ user.role }}')
    doc.add_paragraph('Подразделение: {{ user.unit }}')
    doc.add_paragraph('Дата начала: {{ date_from }}')
    doc.add_paragraph('Дата окончания: {{ date_to }}')
    doc.add_paragraph('Дата запроса: {{ date_req }}')

    template_bytes = io.BytesIO()
    doc.save(template_bytes)
    template_bytes.seek(0)
    return template_bytes.getvalue()


def test_generate_doc_returns_bytesio():
    """Тест: функция возвращает BytesIO объект."""
    tpl_content = create_test_template()
    with open('docs_template/заявление_на_отпуск_бс.docx', 'rb') as rdr:
        tpl_content = rdr.read()
    context = {
        'status': 'ok',
        'user': {
            'name': 'Иван',
            'name_gen': 'Ивана',
            'surname': 'Петров',
            'surname_gen': 'Петрова',
            'father_name': 'Сергеевич',
            'father_name_gen': 'Сергеевича',
            'role': 'Инженер',
            'unit': 'Отдел разработки',
        },
        'date_from': '01.01.2024',
        'date_to': '31.12.2024',
        'date_req': '15.03.2024'
    }

    result = generate_doc(tpl_content, context)
    with open('docs_template/заявление_на_отпуск_бс_333.docx', 'wb') as wr:
        wr.write(result.getvalue())
    assert isinstance(result, io.BytesIO)


def test_generate_doc_metadata():
    """Тест: проверка установки метаданных документа."""
    tpl_content = create_test_template()
    context = {
        'status': 'ok',
        'user': {
            'name': 'Иван',
            'name_gen': 'Ивана',
            'surname': 'Петров',
            'surname_gen': 'Петрова',
            'father_name': 'Сергеевич',
            'father_name_gen': 'Сергеевича',
            'role': 'Инженер',
            'unit': 'Отдел разработки',
        },
        'date_from': '01.01.2024',
        'date_to': '31.12.2024',
        'date_req': '15.03.2024'
    }

    result = generate_doc(tpl_content, context)

    # Читаем полученный документ и проверяем метаданные
    doc = Document(result)
    assert doc.core_properties.author == 'ЗАО СММ'
    assert doc.core_properties.title == 'Документ'
    assert doc.core_properties.subject == 'Документ'
    assert doc.core_properties.keywords == 'документ'


def test_generate_doc_content():
    """Тест: проверка корректности заполнения шаблона."""
    tpl_content = create_test_template()
    context = {
        'status': 'ok',
        'user': {
            'name': 'Иван',
            'name_gen': 'Ивана',
            'surname': 'Петров',
            'surname_gen': 'Петрова',
            'father_name': 'Сергеевич',
            'father_name_gen': 'Сергеевича',
            'role': 'Инженер',
            'unit': 'Отдел разработки',
        },
        'date_from': '01.01.2024',
        'date_to': '31.12.2024',
        'date_req': '15.03.2024'
    }

    result = generate_doc(tpl_content, context)

    # Проверяем, что данные подставились правильно
    doc = Document(result)
    paragraphs = [p.text for p in doc.paragraphs]

    expected_paragraphs = [
        'Статус: ok',
        'Имя: Иван',
        'Фамилия: Петров',
        'Отчество: Сергеевич',
        'Роль: Инженер',
        'Подразделение: Отдел разработки',
        'Дата начала: 01.01.2024',
        'Дата окончания: 31.12.2024',
        'Дата запроса: 15.03.2024'
    ]

    assert paragraphs == expected_paragraphs


def test_generate_doc_with_different_data():
    """Тест: проверка с другими данными."""
    tpl_content = create_test_template()
    context = {
        'status': 'pending',
        'user': {
            'name': 'Петр',
            'name_gen': 'Петра',
            'surname': 'Иванов',
            'surname_gen': 'Иванова',
            'father_name': 'Иванович',
            'father_name_gen': 'Ивановича',
            'role': 'Менеджер',
            'unit': 'Отдел продаж',
        },
        'date_from': '01.02.2024',
        'date_to': '28.02.2024',
        'date_req': '20.01.2024'
    }

    result = generate_doc(tpl_content, context)
    doc = Document(result)
    paragraphs = [p.text for p in doc.paragraphs]

    expected_paragraphs = [
        'Статус: pending',
        'Имя: Петр',
        'Фамилия: Иванов',
        'Отчество: Иванович',
        'Роль: Менеджер',
        'Подразделение: Отдел продаж',
        'Дата начала: 01.02.2024',
        'Дата окончания: 28.02.2024',
        'Дата запроса: 20.01.2024'
    ]

    assert paragraphs == expected_paragraphs


def test_generate_doc_with_empty_values():
    """Тест: проверка с пустыми значениями."""
    tpl_content = create_test_template()
    context = {
        'status': '',
        'user': {
            'name': '',
            'name_gen': '',
            'surname': '',
            'surname_gen': '',
            'father_name': '',
            'father_name_gen': '',
            'role': '',
            'unit': '',
        },
        'date_from': '',
        'date_to': '',
        'date_req': ''
    }

    result = generate_doc(tpl_content, context)
    doc = Document(result)
    paragraphs = [p.text for p in doc.paragraphs]

    expected_paragraphs = [
        'Статус: ',
        'Имя: ',
        'Фамилия: ',
        'Отчество: ',
        'Роль: ',
        'Подразделение: ',
        'Дата начала: ',
        'Дата окончания: ',
        'Дата запроса: '
    ]

    assert paragraphs == expected_paragraphs